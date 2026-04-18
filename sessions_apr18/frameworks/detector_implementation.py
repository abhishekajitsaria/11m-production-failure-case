#!/usr/bin/env python3
"""
WOS MANDATORY PROMPTS v3.4 — SINGLE RUNNABLE SCRIPT
Run against any production DOCX file.
Usage: python3 WOS_PROMPTS_v3_4_TEST.py FILENAME.docx

v3.4 patches four detector gaps surfaced in S27 audit of POST v4_23 /
PRE v4_13 / REWRITE v4_18:

  D1 (mid-word truncations): relaxed short-tail threshold from 1-3 chars
      to 1-4 chars AND added single-char-tail catch (missed POST p928
      '...Rule MV-R2 a' and p1186 '...terminal value from a' in v3.3).

  D7 (mid-list truncations): added list-completeness check. If an
      enumerated list reaches its stated item count (e.g., "ALL TEN
      CRITERIA" + 10 items present) and ends cleanly on content, do
      not flag absence of terminal punctuation. False-positive on
      MASTER Screener ten-criteria list in v3.3.

  D10 (rule coverage): added parenthetical-form definition matching.
      Regex now catches '(R10 — War Dip Protocol)' body-text formal
      definitions, not only paragraph-head 'R10:' forms. REWRITE R10
      missed in v3.3 census because defs are parenthetical.

  D11 (numerical consistency): added calculation-context guard. ₹ amounts
      inside 'SIP × years @ CAGR' arithmetic clauses are calculation
      outcomes, not portfolio-total identity claims. POST p1204 false-
      positive (₹1.8 Crore = 20-year SIP outcome, not portfolio drift).

STRUCTURAL INTEGRITY (from v3.2):
  D1: Mid-word truncations  [PATCHED v3.4]
  D2: Editorial notes-to-self leaked into production
  D3: Numeric placeholders and fill-me-in markers
  D4: Walls of text (Ghost v8 Dimension 7)
  D5: Orphan headings and misplaced structural markers
  D6: Duplicate paragraph content
  D7: Mid-list truncations  [PATCHED v3.4]
  D8: Bracketed placeholders

CONTENT CONSISTENCY (v3.3 additions):
  D9:  Adjacent paragraph repetition (R7 triplet mechanical component)
  D10: Rule/Protocol coverage  [PATCHED v3.4]
  D11: Numerical consistency  [PATCHED v3.4]
  D12: Rupee/Dollar pairing (every ₹X should have $X equivalent)
  D13: Date consistency (same event, different date strings)
  D14: Heading numbering continuity (chapter/sub-chapter gaps)
  D15: Part endgame positioning (endgame at end of Part, not start)
  D16: Sub-chapter placement (34A after 34, not after 35)
  D17: TOC-vs-body consistency (TOC chapter count vs body heading count)
  D18: Rule/Appendix/Protocol cross-references (beyond just Chapter refs)

Each detector reports: count, severity (FAIL/WARN/INFO), sample indices.
Calibrated to exclude workbook sections, CAG conventions, TYPESETTER blocks.
Remaining 15% of manuscript quality requires Rajiv Ghost v8 human reading.
"""

import sys, re, zipfile, hashlib, os
from docx import Document

if len(sys.argv) < 2:
    print("Usage: python3 WOS_PROMPTS_v3_4_TEST.py FILENAME.docx")
    sys.exit(1)

fname = sys.argv[1]
doc = Document(fname)
paras = doc.paragraphs
text = '\n'.join([p.text for p in paras])
words = sum(len(p.text.split()) for p in paras if p.text.strip())

with zipfile.ZipFile(fname) as z:
    xml = z.read('word/document.xml').decode('utf-8')
    zip_names = z.namelist()

with open(fname, 'rb') as f:
    sha = hashlib.sha256(f.read()).hexdigest()[:8]

fails = []
warnings = []

# ============================================================
# PROMPT 1: QC 22-POINT
# ============================================================
print(f"{'='*60}")
print(f"PROMPT 1: QC 22-POINT — {os.path.basename(fname)}")
print(f"{'='*60}")

# QC 1-5: Structure
h1 = sum(1 for p in paras if p.style and p.style.name == 'Heading 1')
h2 = sum(1 for p in paras if p.style and p.style.name == 'Heading 2')
h3 = sum(1 for p in paras if p.style and p.style.name == 'Heading 3')
print(f"1.  Words: {words:,}")
print(f"2.  Paras: {len(paras)}")
print(f"3.  H1: {h1}")
print(f"4.  H2: {h2}")
print(f"5.  H3: {h3}")

# QC 6: Normal
n = sum(1 for p in paras if p.style and p.style.name == 'Normal' and p.text.strip())
r = 'PASS' if n == 0 else 'FAIL'
if n > 0: fails.append(f"QC6 Normal={n}")
print(f"6.  Normal: {n} {r}")

# QC 7: Tables
print(f"7.  Tables: {len(doc.tables)}")

# QC 8-9: Tracked changes
d = len(re.findall(r'w:del\b', xml))
i = len(re.findall(r'w:ins\b', xml))
if d > 0: fails.append(f"QC8 w:del={d}")
if i > 0: fails.append(f"QC9 w:ins={i}")
print(f"8.  w:del: {d} {'PASS' if d==0 else 'FAIL'}")
print(f"9.  w:ins: {i} {'PASS' if i==0 else 'FAIL'}")

# QC 10: Stutters
st = sum(1 for p in paras if re.search(r'\bthe\s+the\b', p.text, re.I))
if st > 0: fails.append(f"QC10 stutters={st}")
print(f"10. Stutters: {st} {'PASS' if st==0 else 'FAIL'}")

# QC 11-13: Frequency
for word, limit, qcnum in [('documented', 120, 11), ('architecture', 120, 12), ('march 23', 20, 13)]:
    c = text.lower().count(word)
    ok = c <= limit
    if not ok: fails.append(f"QC{qcnum} '{word}'={c}/{limit}")
    print(f"{qcnum}. '{word}': {c}/{limit} {'PASS' if ok else 'FAIL'}")

# QC 14: Blanks (corrected: exclude MV-R4 template)
blanks_raw = re.findall(r'₹\s*\[|\$\s*\[|\[\s*\]|XX%|TBD|TODO', text)
template_count = text.count('₹[Y]')
real_blanks = len(blanks_raw) - template_count
if real_blanks > 0: warnings.append(f"QC14 blanks={real_blanks}")
print(f"14. Blanks: {real_blanks} real ({len(blanks_raw)} raw - {template_count} template) {'PASS' if real_blanks==0 else 'CHECK'}")

# QC 15: Junk
jk = len(re.findall(r'[a-z]{30,}', text))
if jk >= 3: warnings.append(f"QC15 junk={jk}")
print(f"15. Junk: {jk} {'PASS' if jk<3 else 'CHECK'}")

# QC 16: Duplication
dp = text.count('$7.5 million')
if dp > 1: fails.append(f"QC16 dup={dp}")
print(f"16. Dup $7.5M: {dp}x {'PASS' if dp<=1 else 'FAIL'}")

# QC 17: SHA
print(f"17. SHA: {sha}")

# QC 18: Comments
cm = len(re.findall(r'w:commentReference', xml))
orphan_comments = [n for n in zip_names if 'comment' in n.lower()]
if cm > 0: warnings.append(f"QC18 comments={cm}")
print(f"18. Comments: {cm} (orphan files: {len(orphan_comments)}) {'PASS' if cm==0 else 'CHECK'}")

# QC 19-22
bk = len(re.findall(r'w:bookmarkStart', xml))
rs = len(set(re.findall(r'w:rsidR="([A-Fa-f0-9]+)"', xml)))
sc = len(re.findall(r'w:sectPr', xml))
fs = os.path.getsize(fname)
print(f"19. Bookmarks: {bk}")
print(f"20. rsidR: {rs} unique")
print(f"21. Sections: {sc}")
print(f"22. Size: {fs:,}B")

qc_result = "22/22 PASS" if not fails else f"{22-len(fails)}/22 — FAIL: {', '.join(fails)}"
print(f"\nQC RESULT: {qc_result}")

# ============================================================
# PROMPT 2: GHOST v8.0 — Computational passes (4-6)
# ============================================================
print(f"\n{'='*60}")
print(f"PROMPT 2: GHOST v8.0 — Computational")
print(f"{'='*60}")

# Pass 4: Jargon (corrected: word boundaries, hyphen-flexible, body-text only)
body_start = text.find('CHAPTER 0')
if body_start == -1:
    body_start = 0
body = text[body_start:]

terms = {
    'CAGR': r'Compound Annual Growth Rate',
    'XIRR': r'Extended Internal Rate of Return',
    'SIP': r'Systematic Investment Plan',
    'HUF': r'Hindu Undivided Family',
    'SEBI': r'Securities and Exchange Board',
    'LTCG': r'Long.?Term Capital Gain',
    'STCG': r'Short.?Term Capital Gain',
    'ELSS': r'Equity Linked Saving',
    'PPF': r'Public Provident Fund',
    'NPS': r'National Pension',
    'ROE': r'Return on Equity',
    'ROCE': r'Return on Capital Employed',
    'AIF': r'Alternative Investment Fund',
    'BER': r'Base Expense Ratio',
    'UCC': r'Uniform Civil Code',
    'BNPL': r'Buy Now.*Pay Later',
    'VFO': r'Virtual Family Office',
    'ISA': r'Individual Savings Account',
}

jargon_defined = 0
jargon_undefined = []
jargon_total = 0

for term, exp in terms.items():
    m = re.search(rf'\b{term}\b', body)
    if not m:
        continue
    jargon_total += 1
    before = body[:m.start()]
    near = body[m.start():m.start()+1000]
    # Check: expansion before, expansion near, or inverse parenthetical "(TERM)" before
    ok = (re.search(exp, before, re.I) or
          re.search(exp, near, re.I) or
          re.search(rf'\({term}\)', before))
    if ok:
        jargon_defined += 1
    else:
        jargon_undefined.append(term)

print(f"\nPass 4 Jargon: {jargon_defined}/{jargon_total} defined at first body use")
if jargon_undefined:
    print(f"  Undefined: {', '.join(jargon_undefined)}")
    print(f"  Note: SEBI/SIP/HUF/CAGR are household terms in Indian finance")

# Pass 5: Signpost
wtm = len(re.findall(r'WHY THIS MATTERS', text))
od = len(re.findall(r'ONE DECISION', text))
print(f"\nPass 5 Signpost: {wtm} WHY THIS MATTERS, {od} ONE DECISION")

# Pass 5B: Per-chapter ONE DECISION coverage
chapter_starts = []
for idx, p in enumerate(paras):
    if p.style and 'Heading' in p.style.name and re.match(r'CHAPTER \d+', p.text):
        chapter_starts.append((idx, p.text.split(':')[0].strip() if ':' in p.text else p.text[:30]))
chapters_missing_od = []
for ci, (start, title) in enumerate(chapter_starts):
    end = chapter_starts[ci+1][0] if ci+1 < len(chapter_starts) else len(paras)
    ch_text = ' '.join(paras[j].text for j in range(start, end))
    if 'ONE DECISION' not in ch_text:
        chapters_missing_od.append(title)
if chapters_missing_od:
    print(f"  Chapters WITHOUT ONE DECISION: {', '.join(chapters_missing_od)}")
else:
    print(f"  All {len(chapter_starts)} chapters have ONE DECISION")

# Pass 5C: Part endgame closures
part_headings = []
for idx, p in enumerate(paras):
    if p.style and p.style.name == 'Heading 1':
        m = re.match(r'PART ([A-D])', p.text)
        if m:
            part_headings.append((idx, m.group(1), p.text.strip()))
# Deduplicate: only keep first occurrence of each Part letter
seen_parts = {}
unique_parts = []
for idx, letter, txt in part_headings:
    if letter not in seen_parts:
        seen_parts[letter] = (idx, letter, txt)
        unique_parts.append((idx, letter, txt))
endgame_keywords = ['you now have', 'you have now', 'your family', 'you can now evaluate', 'is now structured', 'is now in place', 'complete operating system', 'governance is now', 'financial structure is now']
parts_missing_endgame = []
for pi in range(1, len(unique_parts)):
    part_idx = unique_parts[pi][0]
    part_letter = unique_parts[pi][1]
    prev_letter = unique_parts[pi-1][1]
    search_range = ' '.join(paras[j].text.lower() for j in range(max(0, part_idx-5), part_idx))
    has_endgame = any(kw in search_range for kw in endgame_keywords)
    if not has_endgame:
        parts_missing_endgame.append(f"PART {prev_letter} (before PART {part_letter})")
if parts_missing_endgame:
    print(f"  Parts WITHOUT endgame closure: {', '.join(parts_missing_endgame)}")
else:
    print(f"  All Part transitions have endgame closures")

# Pass 5D: Quick Install achievement statement
qi_text = ''
for idx, p in enumerate(paras):
    if 'quick install' in p.text.lower() and p.style and 'Heading' in p.style.name:
        qi_text = ' '.join(paras[j].text for j in range(idx, min(len(paras), idx+10)))
        break
has_qi_achieve = 'WHAT YOU WILL ACHIEVE' in qi_text or 'what you will achieve' in qi_text.lower()
print(f"  Quick Install achievement: {'PASS' if has_qi_achieve else 'MISSING'}")

# Pass 6: Structural (covered by QC above)
print(f"\nPass 6 Structural: {'PASS' if not fails else 'FAIL — see QC'}")

# ============================================================
# PROMPT 3: PRE-PRINTING CHECKLIST
# ============================================================
print(f"\n{'='*60}")
print(f"PROMPT 3: PRE-PRINTING CHECKLIST")
print(f"{'='*60}")

# 3A: Typesetter placeholders
ts_blocks = [p.text for p in paras if 'TYPESETTER' in p.text]
print(f"\n3A Typesetter: {len(ts_blocks)} blocks")
for t in ts_blocks:
    print(f"    {t[:80]}...")

# 3B: Cross-reference integrity
ch_headings = set()
for p in paras:
    m = re.match(r'CHAPTER (\d+[A-B]?)', p.text)
    if m and p.style and 'Heading' in p.style.name:
        ch_headings.add(m.group(1))
xrefs = set(re.findall(r'Chapter (\d+[A-B]?)', text))
broken = sorted(
    [x for x in xrefs if x not in ch_headings
     and re.match(r'\d+', x) and int(re.match(r'\d+', x).group()) <= 49],
    key=lambda x: int(re.match(r'\d+', x).group())
)
print(f"\n3B Cross-refs: {len(ch_headings)} headings, {len(xrefs)} references, {len(broken)} broken")
if broken:
    print(f"    Broken: {', '.join(broken)}")

# 3C: Chapter numbering
print(f"\n3C Chapters in headings: {sorted(ch_headings, key=lambda x: (int(re.match(r'(\\d+)', x).group()), x) if re.match(r'\\d+', x) else (999, x))[:10]}...")
subs_expected = ['12A', '21A', '28A', '28B', '32A', '32B', '34A', '43A']
subs_found = [s for s in subs_expected if s in ch_headings]
subs_missing = [s for s in subs_expected if s not in ch_headings]
print(f"    Sub-chapters: {len(subs_found)}/8 found, missing: {', '.join(subs_missing) if subs_missing else 'none'}")

# 3D: Front matter order
print(f"\n3D Front matter:")
h1_list = [(j, p.text[:60]) for j, p in enumerate(paras) if p.style and p.style.name == 'Heading 1']
for j, t in h1_list[:8]:
    print(f"    p{j:4d}: {t}")

# 3E: Back matter order
print(f"\n3E Back matter:")
for j, t in h1_list[-6:]:
    print(f"    p{j:4d}: {t}")

# 3F: Copyright page
print(f"\n3F Copyright:")
cp_checks = [
    ('©', 'Copyright ©'),
    ('First published', 'First published'),
    ('ISBN', 'ISBN'),
    ('AI TRANSPARENCY', 'AI Transparency Disclosure'),
    ('not a SEBI-registered', 'SEBI not-registered disclaimer'),
]
cp_pass = 0
for term, label in cp_checks:
    found = term in text
    if found: cp_pass += 1
    print(f"    {'PASS' if found else 'MISSING'} {label}")

# 3G: Hyphenation (corrected: exclude fund names)
print(f"\n3G Hyphenation:")
lt_all = len(re.findall(r'(?<!-)long term(?!-)', text, re.I))
lt_fund = text.count('Long Term Equity')
lt_real = lt_all - lt_fund
st_real = len(re.findall(r'(?<!-)short term(?!-)', text, re.I))
print(f"    long-term: {lt_real} unhyphenated ({lt_fund} fund names excluded) {'PASS' if lt_real==0 else 'FIX'}")
print(f"    short-term: {st_real} unhyphenated {'PASS' if st_real==0 else 'FIX'}")

# 3H: Proofreading
print(f"\n3H Proofreading:")
dbl_spaces = len(re.findall(r'  ', text))
smart_quotes_mixed = bool(re.search(r"['\"].*['\u2018\u2019\u201c\u201d]", text[:5000]))
print(f"    Double spaces: {dbl_spaces} {'PASS' if dbl_spaces==0 else 'FIX'}")

# 3I: Data verification
print(f"\n3I Data points:")
data_checks = [
    ('₹7.8 Crore', 'Portfolio total'),
    ('₹2,885.44', 'BSE avg cost'),
    ('₹3,86,648', 'LTCG carry-forward'),
]
for term, label in data_checks:
    found = term in text
    print(f"    {'PASS' if found else 'CHECK'} {label} ({term})")

# 3J: Page estimate
pages = int(words / 250)
print(f"\n3J Page estimate: ~{pages} pages ({words:,}w at 250w/page)")

# ============================================================
# PROMPT 4: CONTRA DEFENSE + INTEGRATION MAP
# ============================================================
print(f"\n{'='*60}")
print(f"PROMPT 4: CONTRA DEFENSE + INTEGRATION MAP")
print(f"{'='*60}")

contra_pass = 0
contra_total = 0

# 4A: SEBI media exemption
contra_total += 1
has_sebi_media = '2(1)(l)' in text or 'media exemption' in text.lower()
print(f"\n4A SEBI media exemption: {'PASS' if has_sebi_media else 'MISSING'}")
if has_sebi_media: contra_pass += 1

# 4B: Architecture-vs-information defense (Varsity/free content)
contra_total += 1
has_arch_defense = 'varsity' in text.lower() and ('architecture' in text.lower() or 'structure' in text.lower())
print(f"4B Free content defense (Varsity): {'PASS' if has_arch_defense else 'MISSING'}")
if has_arch_defense: contra_pass += 1

# 4C: Past performance disclaimer
contra_total += 1
has_past_perf = 'past performance' in text.lower() or 'past returns' in text.lower() or 'historical' in text.lower()
print(f"4C Past performance disclaimer: {'PASS' if has_past_perf else 'MISSING'}")
if has_past_perf: contra_pass += 1

# 4D: Reader A ₹500 SIP content
contra_total += 1
has_reader_a = 'reader a' in text.lower() and '500' in text
print(f"4D Reader A ₹500 path: {'PASS' if has_reader_a else 'MISSING'}")
if has_reader_a: contra_pass += 1

# 4E: Privacy — patriarch title scrubbed
contra_total += 1
patriarch_exposed = 'President of the Gauhati Stock Exchange' in text
print(f"4E Privacy patriarch: {'PASS (scrubbed)' if not patriarch_exposed else 'EXPOSED'}")
if not patriarch_exposed: contra_pass += 1

# 4F: Privacy — minor formal name scrubbed
contra_total += 1
minor_exposed = 'Bhavamanyu' in text
print(f"4F Privacy minor: {'PASS (scrubbed)' if not minor_exposed else 'EXPOSED'}")
if not minor_exposed: contra_pass += 1

# 4G: Methodology-over-output framing
contra_total += 1
has_method = 'methodology over output' in text.lower() or 'stocks are evidence' in text.lower() or 'system is what you are building' in text.lower() or 'value of this book is in the rules' in text.lower()
print(f"4G Methodology framing: {'PASS' if has_method else 'MISSING'}")
if has_method: contra_pass += 1

# 4H: Integration map critical data points
print(f"\n4H Integration map data points:")
map_checks = [
    ('68%', 'retirement', '68% no retirement plan'),
    ('72%', 'goal', '72% no documented goal'),
    ('378', 'millionaire', '378K Indian millionaires'),
    ('225 million', 'demat', '225M demat gap'),
    ('23%', 'gold', 'Gold 23% household'),
    ('28%', 'fixed deposit', 'FD 28% household'),
    ('4.2', 'debt', 'Avg debt ₹4.2L'),
    ('73%', 'QR', 'QR 73% adoption'),
    ('Varsity', '17', 'Varsity 17 modules'),
    ('2(1)(l)', '', 'SEBI media exemption'),
]
map_present = 0
for marker, ctx, desc in map_checks:
    found = any(marker in p.text and (not ctx or ctx.lower() in p.text.lower()) for p in paras)
    if found: map_present += 1
    print(f"    {'PASS' if found else 'MISSING'} {desc}")

contra_result = f"{contra_pass}/{contra_total} contra, {map_present}/10 integration map"

# ============================================================
# PROMPT 5: DETECTORS ENGINE (v3.2 addition)
# Eight detectors for manuscript-integrity issues that structural checks miss
# ============================================================
print(f"\n{'='*60}")
print(f"PROMPT 5: DETECTORS ENGINE")
print(f"{'='*60}")

detector_fails = []
detector_warns = []

# Shared helpers
COMPLETE_SHORT = {
    'a', 'an', 'as', 'at', 'be', 'by', 'do', 'go', 'he', 'if', 'in', 'is', 'it',
    'me', 'my', 'no', 'of', 'on', 'or', 'so', 'to', 'up', 'us', 'we', 'are',
    'all', 'any', 'but', 'can', 'did', 'for', 'had', 'has', 'her', 'him', 'his',
    'how', 'its', 'may', 'not', 'now', 'our', 'out', 'she', 'the', 'too', 'two',
    'was', 'who', 'why', 'you', 'yet', 'few', 'let', 'off', 'per', 'set', 'top',
    'way', 'win', 'own', 'run', 'use', 'low', 'new', 'old', 'one', 'see', 'end',
    'get', 'big', 'say', 'yes'
}
TERMINAL_PUNCT = ('.', '!', '?', '"', ')', ']', ':', ';', '*', '%', '—', '"', ''', ''')

# ------------------------------------------------------------
# D1: Mid-word truncations  [v3.4 PATCHED]
# Detects paragraphs ending in short tails that indicate incomplete content.
# v3.3 used 1-3 char tails only, missing POST p928 '...Rule MV-R2 a' and
# p1186 '...terminal value from a'. v3.4 adds a secondary pattern for
# paragraphs ending on an article/preposition (a/an/the/to/from/of etc.)
# — these cannot legitimately end a declarative sentence, so their
# appearance at paragraph end with no terminal punct indicates truncation.
# Case-sensitive to avoid flagging title-case headers like "Who This Book Is For".
# ------------------------------------------------------------
d1_hits = []
# Articles and prepositions that cannot end a declarative sentence.
# Match case-sensitive lowercase only (title-case "For" in headers is legitimate).
D1_TRUNC_WORDS = {
    'a', 'an', 'the', 'to', 'from', 'of', 'with', 'for',
    'by', 'on', 'in', 'at', 'as', 'is', 'and', 'or', 'but',
    'into', 'onto', 'upon', 'than', 'that', 'this', 'these', 'those',
}
for idx, p in enumerate(paras):
    t = p.text.rstrip()
    if not t or len(t) < 20:
        continue
    if t.endswith(TERMINAL_PUNCT):
        continue
    if re.search(r'\d$', t):
        continue
    if len(t.split()) < 3 and t.isupper():
        continue
    # Primary: 1-3 char lowercase tail (from v3.3), filtered by COMPLETE_SHORT
    m = re.search(r' ([a-z]{1,3})$', t)
    if m and m.group(1) not in COMPLETE_SHORT:
        d1_hits.append((idx, m.group(1), t[-50:]))
        continue
    # Secondary [v3.4]: paragraph ends on a lowercase article/preposition.
    # These cannot legitimately end a declarative sentence. Case-sensitive:
    # title-case "For" in headers (e.g., "Who This Book Is For") is legitimate.
    # Catches POST p928 '...Rule MV-R2 a' and p1186 '...terminal value from a'.
    m2 = re.search(r' ([a-z]+)$', t)
    if m2 and m2.group(1) in D1_TRUNC_WORDS:
        d1_hits.append((idx, m2.group(1), t[-50:]))

d1_status = 'PASS' if len(d1_hits) == 0 else 'FAIL'
if d1_hits:
    detector_fails.append(f"D1 truncations={len(d1_hits)}")
print(f"\nD1 Mid-word truncations: {len(d1_hits)} {d1_status}")
for idx, frag, tail in d1_hits[:10]:
    print(f"    [p{idx}] ends '...{frag}' — ...{tail}")
if len(d1_hits) > 10:
    print(f"    ... ({len(d1_hits) - 10} more)")

# ------------------------------------------------------------
# D2: Editorial notes-to-self leaked into production
# Catches author/editor directives that shouldn't be in production file:
#   - "should remain — companies are public entities"
#   - "should be preserved with retrospective framing"
#   - "fair game for analysis per SEBI"
#   - "TO DO", "FIXME", "NOTE TO EDITOR"
# Excludes legitimate TYPESETTER blocks (already counted in 3A)
# ------------------------------------------------------------
editorial_patterns = [
    (r'should remain\s*[—–-]\s*companies', 'editor directive (company refs)'),
    (r'should be preserved with retrospective', 'editor directive (preservation)'),
    (r'fair game for analysis per SEBI', 'editor directive (SEBI rationale)'),
    (r'\bTO\s*DO\b', 'TODO marker'),
    (r'\bFIXME\b', 'FIXME marker'),
    (r'NOTE TO (EDITOR|AUTHOR|SELF)', 'editor note'),
    (r'NEEDS (AUTHOR|EDITOR|MORE) REVIEW', 'review flag'),
    (r'<!--\s*[A-Z]', 'HTML comment'),
]
d2_hits = []
for idx, p in enumerate(paras):
    # Skip typesetter blocks (legitimate)
    if 'TYPESETTER' in p.text:
        continue
    for pat, label in editorial_patterns:
        if re.search(pat, p.text):
            d2_hits.append((idx, label, p.text[:80]))
            break

d2_status = 'PASS' if len(d2_hits) == 0 else 'FAIL'
if d2_hits:
    detector_fails.append(f"D2 editorial_notes={len(d2_hits)}")
print(f"\nD2 Editorial notes in production: {len(d2_hits)} {d2_status}")
for idx, label, text in d2_hits[:10]:
    print(f"    [p{idx}] {label}: {text}")
if len(d2_hits) > 10:
    print(f"    ... ({len(d2_hits) - 10} more)")

# ------------------------------------------------------------
# D3: Numeric placeholders and fill-me-in markers
# ₹[X], $[Y], [TBD], [XX], [YY], PLACEHOLDER
# CALIBRATION: Excludes workbook/worksheet fill-in blanks (intentional for reader use)
# A workbook paragraph typically has label + colon + blank pattern: "Total: ₹___"
# ------------------------------------------------------------
placeholder_patterns = [
    (r'₹\s*\[(?!Y\])[A-Z]\]', 'rupee placeholder'),
    (r'\$\s*\[(?!Y\])[A-Z]\]', 'dollar placeholder'),
    (r'\[TBD\]', 'TBD'),
    (r'\[XX+\]', 'XX placeholder'),
    (r'\[YY+\]', 'YY placeholder'),
    (r'\bPLACEHOLDER\b', 'PLACEHOLDER'),
]
# Find the workbook / prepress checklist section boundary (back matter)
# Typically appendix-level. Use multiple markers since conventions vary.
workbook_start = None
workbook_markers = (
    'COMPANION WORKBOOK',
    'PREPRESS CHECKLIST',
    'SESSION WORKSHEET',
    'WORKSHEET 1',  # WOS-specific: worksheet series
    'THE SESSION WORKBOOK',
    'IMPLEMENTATION WORKSHEET',
)
for idx, p in enumerate(paras):
    t = p.text.upper().strip()
    if any(marker in t for marker in workbook_markers):
        workbook_start = idx
        break

d3_hits = []
for idx, p in enumerate(paras):
    # Skip workbook/worksheet section — blanks are intentional there
    if workbook_start is not None and idx >= workbook_start:
        continue
    for pat, label in placeholder_patterns:
        m = re.search(pat, p.text)
        if m:
            d3_hits.append((idx, label, p.text[max(0, m.start()-20):m.end()+20]))
            break

d3_status = 'PASS' if len(d3_hits) == 0 else 'FAIL'
if d3_hits:
    detector_fails.append(f"D3 placeholders={len(d3_hits)}")
print(f"\nD3 Numeric/fill-in placeholders: {len(d3_hits)} {d3_status}")
for idx, label, snippet in d3_hits[:10]:
    print(f"    [p{idx}] {label}: ...{snippet}...")

# ------------------------------------------------------------
# D4: Walls of text (Ghost v8 Dimension 7 — visual pace)
# Paragraphs with >150 words are dense blocks that damage page-30 test
# WARN at 150-200w, FAIL at 250+w (hard-to-read)
# ------------------------------------------------------------
walls = []
for idx, p in enumerate(paras):
    # Skip typesetter blocks and obvious list-pastes
    if 'TYPESETTER' in p.text:
        continue
    wc = len(p.text.split())
    if wc >= 150:
        walls.append((idx, wc, p.text[:60]))

hard_walls = [w for w in walls if w[1] >= 250]
soft_walls = [w for w in walls if 150 <= w[1] < 250]
d4_status = 'PASS' if len(hard_walls) == 0 else 'WARN' if len(hard_walls) < 5 else 'FAIL'
if len(hard_walls) >= 5:
    detector_fails.append(f"D4 hard_walls={len(hard_walls)}")
elif len(hard_walls) > 0 or len(soft_walls) >= 20:
    detector_warns.append(f"D4 walls={len(hard_walls)}/{len(soft_walls)}")
print(f"\nD4 Walls of text: {len(hard_walls)} hard (≥250w), {len(soft_walls)} soft (150-249w) {d4_status}")
for idx, wc, text in hard_walls[:5]:
    print(f"    [p{idx}] {wc}w: {text}...")

# ------------------------------------------------------------
# D5: Orphan headings and misplaced structural markers
# Heading-styled paragraph followed by non-body or another heading (no content)
# Also: body paragraph that looks like a heading ("CHAPTER X" as Body Text)
# CALIBRATION: Excludes "CHAPTER AT A GLANCE" — project convention is Body Text
# Excludes "CHAPTER N: WHY THIS MATTERS" — signpost convention
# ------------------------------------------------------------
CAG_MARKERS = {'CHAPTER AT A GLANCE'}
d5_hits = []
for idx, p in enumerate(paras):
    style = p.style.name if p.style else ''
    t = p.text.strip()
    # Body-text paragraphs that look like headings — BUT not CAG markers (project convention)
    if 'Heading' not in style and 'Title' not in style:
        if t in CAG_MARKERS:
            continue
        if re.match(r'^(CHAPTER|PART|APPENDIX)\s+[A-Z0-9]', t) and len(t.split()) <= 6:
            # Also skip "CHAPTER N: WHY THIS MATTERS" signposts
            if 'WHY THIS MATTERS' in t:
                continue
            d5_hits.append((idx, 'body-text-heading', f'Looks like heading but styled {style}: {t}'))
    # Heading followed by another heading (empty content between)
    if 'Heading' in style:
        if idx + 1 < len(paras):
            next_p = paras[idx + 1]
            next_style = next_p.style.name if next_p.style else ''
            next_t = next_p.text.strip()
            # Skip the known sub-heading stacking (Heading 1→Heading 2, Heading 2→Heading 3)
            if 'Heading' in next_style and not next_t:
                d5_hits.append((idx, 'empty-heading', f'{style}: {t} → empty follow'))

d5_status = 'PASS' if len(d5_hits) == 0 else 'WARN'
if d5_hits:
    detector_warns.append(f"D5 orphan_headings={len(d5_hits)}")
print(f"\nD5 Orphan/misplaced headings: {len(d5_hits)} {d5_status}")
for idx, label, snippet in d5_hits[:10]:
    print(f"    [p{idx}] {label}: {snippet}")

# ------------------------------------------------------------
# D6: Duplicate paragraph content (substantial length)
# Body paragraphs ≥30 words appearing more than once in the document
# Excludes: CAG headers, repeated signposts, known templates
# ------------------------------------------------------------
from collections import Counter
seen_texts = Counter()
text_first_idx = {}
for idx, p in enumerate(paras):
    t = p.text.strip()
    wc = len(t.split())
    if wc < 30:
        continue
    # Skip template-like repetitions
    if t.startswith('CHAPTER AT A GLANCE'):
        continue
    if t.startswith('ONE DECISION'):
        continue
    if t.startswith('WHY THIS MATTERS'):
        continue
    if 'TYPESETTER' in t:
        continue
    seen_texts[t] += 1
    if t not in text_first_idx:
        text_first_idx[t] = idx

d6_hits = [(text_first_idx[t], count, t[:80]) for t, count in seen_texts.items() if count > 1]
d6_status = 'PASS' if len(d6_hits) == 0 else 'WARN'
if d6_hits:
    detector_warns.append(f"D6 duplicates={len(d6_hits)}")
print(f"\nD6 Duplicate substantive paragraphs: {len(d6_hits)} {d6_status}")
for idx, count, text in d6_hits[:10]:
    print(f"    [p{idx}] appears {count}x: {text}...")

# ------------------------------------------------------------
# D7: Mid-list truncations (lone numerals, broken bullet sequences)  [v3.4 PATCHED]
# CALIBRATION: Only flags paragraphs that genuinely contain enumerated list
# ("1. X? 2. Y? 3. Z?" pattern) AND end mid-way. Not dates, not labels.
# Catches the Porter-list case: "1. X? 2. Y? 3. Z? 4. W? 5" (item 5 missing)
# v3.4: List-completeness check — if the paragraph declares a count
# ("ALL TEN CRITERIA", "five structural themes", "Three Rules") AND
# the enumerated items reach that count AND the last item has real content,
# do NOT flag. False-positive in v3.3 on MASTER Screener: "ALL TEN CRITERIA
# 1. X 2. Y ... 10. P/FCF: ≤20" (complete list ending without period).
# ------------------------------------------------------------

# Declared-count detection: find patterns like "ALL TEN CRITERIA", "five X",
# "the three Y", "Seven Z" that state the list's expected length.
NUMBER_WORDS = {
    'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5, 'six': 6,
    'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10, 'eleven': 11, 'twelve': 12,
}

def declared_count(text_frag):
    """Return stated item count if paragraph declares one, else None."""
    # Pattern: "ALL TEN CRITERIA", "all five", "three rules", "seven items"
    m = re.search(r'\b(?:ALL\s+)?(\w+)\s+(?:CRITERIA|RULES|ITEMS|STEPS|POINTS|THEMES|QUESTIONS|PRINCIPLES|CONDITIONS|FACTORS|CHECKS|TESTS|GATES|PHASES|PARTS|SIGNALS)\b', text_frag, re.IGNORECASE)
    if m:
        word = m.group(1).lower()
        if word in NUMBER_WORDS:
            return NUMBER_WORDS[word]
        if word.isdigit():
            return int(word)
    return None

d7_hits = []
for idx, p in enumerate(paras):
    t = p.text.rstrip()
    if not t:
        continue
    # Require paragraph to contain multiple numbered list items ("1. X 2. Y" pattern)
    # This excludes paragraphs that just happen to end in "2020." or "Instrument 3."
    list_items = re.findall(r'(?:^|\s)(\d+)\.\s+\S', t)
    if len(list_items) < 3:
        continue  # Not an enumerated list — skip

    # v3.4: List-completeness check. If the paragraph declares a count and
    # the number of enumerated items matches, the list is complete — do not
    # flag even if the paragraph ends without terminal punctuation.
    decl = declared_count(t)
    if decl is not None and len(list_items) >= decl:
        # Extract the last declared item's content and verify it's substantive.
        last_item_text_match = re.search(rf'(?:^|\s){decl}\.\s+(.+?)$', t, re.DOTALL)
        if last_item_text_match:
            last_item_text = last_item_text_match.group(1).strip()
            # Complete list if last item has >= 1 substantive word
            if len(last_item_text.split()) >= 1:
                continue  # List is complete — no flag

    # Pattern A: ends with bare numeral on own line after full list items
    # e.g., "...force price reductions?\n5"
    if re.search(r'\n\d+\s*$', t):
        d7_hits.append((idx, 'lone-numeral-terminator', t[-50:]))
        continue
    # Pattern B: last list item has no content (just "N. " then end)
    last_item_match = re.search(r'(?:^|\s)(\d+)\.\s*$', t)
    if last_item_match:
        # Exclude year-like 4-digit endings (not list items, just dates)
        last_num_b = last_item_match.group(1)
        if re.match(r'^(19|20)\d{2}$', last_num_b):
            continue
        d7_hits.append((idx, f'list-item-{last_num_b}-empty', t[-50:]))
        continue
    # Pattern C: enumerated list but last item is <3 words before end-of-paragraph
    # (indicates broken list) — but only if no declared count was satisfied above
    # Find all "N. content" sequences
    item_contents = re.findall(r'(?:^|\s)(\d+)\.\s+(.+?)(?=\s+\d+\.\s|$)', t, re.DOTALL)
    if len(item_contents) >= 3:
        last_num, last_content = item_contents[-1]
        # Exclude year-like matches (last_num is a 4-digit year ending the paragraph)
        # e.g., "LTCG clock: November 2026." — 2026 isn't a list item
        if re.match(r'^(19|20)\d{2}$', last_num):
            continue
        if len(last_content.split()) < 3:
            d7_hits.append((idx, f'list-item-{last_num}-short', t[-60:]))

d7_status = 'PASS' if len(d7_hits) == 0 else 'FAIL'
if d7_hits:
    detector_fails.append(f"D7 mid_list_trunc={len(d7_hits)}")
print(f"\nD7 Mid-list truncations: {len(d7_hits)} {d7_status}")
for idx, label, snippet in d7_hits[:10]:
    print(f"    [p{idx}] {label}: ...{snippet}")

# ------------------------------------------------------------
# D8: Bracketed placeholders not caught elsewhere
# [INSERT X], {value}, <var>, ___ (3+ underscores)
# CALIBRATION: Excludes TYPESETTER blocks (legitimate) and workbook section
# ------------------------------------------------------------
bracket_patterns = [
    (r'\[INSERT\s+[A-Z]', 'INSERT directive'),
    (r'\{[a-z_]+\}', 'curly placeholder'),
    (r'<[a-z_]+>(?!/)', 'angle placeholder'),
    (r'\bXXX+\b', 'XXX placeholder'),
]
# Underscore blanks only flagged OUTSIDE workbook section
underscore_pattern = (r'(?<![a-z])_{3,}(?![a-z])', 'underscore blank')

d8_hits = []
for idx, p in enumerate(paras):
    if 'TYPESETTER' in p.text:
        continue
    for pat, label in bracket_patterns:
        m = re.search(pat, p.text)
        if m:
            d8_hits.append((idx, label, p.text[max(0, m.start()-15):m.end()+15]))
            break
    # Underscore blanks — only flag outside workbook section
    if workbook_start is None or idx < workbook_start:
        m = re.search(underscore_pattern[0], p.text)
        if m:
            d8_hits.append((idx, underscore_pattern[1], p.text[max(0, m.start()-15):m.end()+15]))

d8_status = 'PASS' if len(d8_hits) == 0 else 'FAIL'
if d8_hits:
    detector_fails.append(f"D8 bracket_placeholders={len(d8_hits)}")
print(f"\nD8 Bracketed placeholders: {len(d8_hits)} {d8_status}")
for idx, label, snippet in d8_hits[:10]:
    print(f"    [p{idx}] {label}: ...{snippet}...")

# ------------------------------------------------------------
# D9: Adjacent paragraph repetition (R7 triplet mechanical component)
# Flags where paragraph N and N+1 share ≥40% non-stopword tokens (repetition signal)
# Excludes: CAG templates, rule definitions (R-numbers repeat intentionally), lists
# ------------------------------------------------------------
STOPWORDS = set('''a an as at be by do go he if in is it me my no of on or so to up us we
    are all any but can did for had has her him his how its may not now our out she the
    too two was who why you yet are and but for nor not that this these those what when
    where which while with from have been were will would could should about above after
    before because been being into more most such than then them very very just only each
    both some same very also such''' .split())

def tokens(text):
    """Tokenize to lowercase non-stopword content words (≥4 chars)."""
    toks = re.findall(r'\b[a-z]{4,}\b', text.lower())
    return set(t for t in toks if t not in STOPWORDS)

d9_hits = []
for idx in range(len(paras) - 1):
    pa, pb = paras[idx], paras[idx + 1]
    ta, tb = pa.text.strip(), pb.text.strip()
    # Skip short paragraphs, headings, templates, typesetter
    wca, wcb = len(ta.split()), len(tb.split())
    if wca < 30 or wcb < 30:
        continue
    styles = (pa.style.name if pa.style else '', pb.style.name if pb.style else '')
    if any('Heading' in s for s in styles):
        continue
    if 'TYPESETTER' in ta or 'TYPESETTER' in tb:
        continue
    if ta.startswith(('CHAPTER AT A GLANCE', 'ONE DECISION', 'WHY THIS MATTERS', 'Score 8–10', 'Score 8-10', 'RULE R', 'PROTOCOL')):
        continue
    if tb.startswith(('CHAPTER AT A GLANCE', 'ONE DECISION', 'WHY THIS MATTERS', 'Score 8–10', 'Score 8-10', 'RULE R', 'PROTOCOL')):
        continue
    toks_a, toks_b = tokens(ta), tokens(tb)
    if len(toks_a) < 8 or len(toks_b) < 8:
        continue
    overlap = toks_a & toks_b
    smaller = min(len(toks_a), len(toks_b))
    ratio = len(overlap) / smaller if smaller else 0
    if ratio >= 0.4:
        d9_hits.append((idx, idx + 1, round(ratio, 2), ta[:60]))

d9_status = 'PASS' if len(d9_hits) == 0 else 'WARN'
if d9_hits:
    detector_warns.append(f"D9 adj_repeat={len(d9_hits)}")
print(f"\nD9 Adjacent paragraph repetition: {len(d9_hits)} {d9_status}")
for a, b, ratio, snippet in d9_hits[:10]:
    print(f"    [p{a}]↔[p{b}] {ratio:.0%} overlap: {snippet}...")
if len(d9_hits) > 10:
    print(f"    ... ({len(d9_hits) - 10} more)")

# ------------------------------------------------------------
# D10: Rule/Protocol coverage — defined vs used  [v3.4 PATCHED]
# Finds R1-R21 (and any R-number) definitions and usages across the document.
# Flags: rules used but never formally defined; rules defined but never used.
# v3.4: Added parenthetical-form matching for definitions like
# "(R10 — War Dip Protocol)" that appear in body text rather than as
# paragraph-head rule registers. v3.3 missed REWRITE R10 (defined at
# p1726 as "(R10 — War Dip Protocol)") and reported it as undefined.
# ------------------------------------------------------------
# Rule definitions: multiple formats — "RULE RN", "Rule RN —", "RN — definition"
rule_defs = set()  # set of rule numbers that have a definition
rule_uses = set()  # set of rule numbers referenced in body
for idx, p in enumerate(paras):
    t = p.text
    # Format A: "RULE R14", "Rule R14", "RULE R14 —"
    for m in re.finditer(r'\bRULE\s+R(\d+[A-Z]?)\b', t):
        rule_defs.add(m.group(1))
    # Format B: "R14 — <content>" at paragraph start (defines the rule inline)
    # Requires dash after the rule number (definition marker)
    m_def = re.match(r'^\s*R(\d+[A-Z]?)\s*[—–-]\s+', t)
    if m_def:
        rule_defs.add(m_def.group(1))
    # Format C: "R14:" at paragraph start with content after
    m_def_c = re.match(r'^\s*R(\d+[A-Z]?)\s*:\s+\S', t)
    if m_def_c:
        rule_defs.add(m_def_c.group(1))
    # Format D [v3.4]: "(R10 — Protocol Name)" or "(R10: Protocol Name)"
    # parenthetical formal definition within body prose.
    for m_paren in re.finditer(r'\(R(\d+[A-Z]?)\s*[—–:-]\s+[A-Z]', t):
        rule_defs.add(m_paren.group(1))
    # Format E [v3.4]: "The R10 <Protocol Name>" — body-text definitional
    # where the rule number + named protocol immediately follow "The".
    # Catches REWRITE p1730 "The R10 War Dip Protocol operationalises..."
    for m_body in re.finditer(r'\bThe\s+R(\d+[A-Z]?)\s+[A-Z][a-z]+\s+(?:Protocol|Rule|Gate|Ceiling|Filter|Cap|Cooling|Block|Limit)', t):
        rule_defs.add(m_body.group(1))
    # Usage pattern: "R10" standalone (not "R10:" since that'd be a definition too)
    for m in re.finditer(r'\bR(\d+[A-Z]?)\b', t):
        rule_uses.add(m.group(1))

# Rules used but never formally defined (uses - defs)
undefined_used = rule_uses - rule_defs
# Rules defined but never referenced (defs - uses, though "defined" means the rule was shown)
# Usually every defined rule has at least its definition reference, so this is usually empty
unused_defined = rule_defs - rule_uses

d10_hits = []
for rn in sorted(undefined_used, key=lambda x: (int(re.match(r'\d+', x).group()), x)):
    d10_hits.append(('undefined-but-used', f'R{rn}'))
for rn in sorted(unused_defined, key=lambda x: (int(re.match(r'\d+', x).group()), x)):
    d10_hits.append(('defined-but-unused', f'R{rn}'))

d10_status = 'PASS' if len(d10_hits) == 0 else 'WARN'
if d10_hits:
    detector_warns.append(f"D10 rule_coverage={len(d10_hits)}")
print(f"\nD10 Rule coverage (defined/used): {len(rule_defs)} defined, {len(rule_uses)} used, {len(d10_hits)} gaps {d10_status}")
for kind, label in d10_hits[:15]:
    print(f"    {kind}: {label}")

# ------------------------------------------------------------
# D11: Numerical consistency — same canonical quantity, different values
# Groups are built from paragraph-local context. Flags: same label → different numbers.
# Conservative: checks a curated list of canonical quantities that recur.
# ------------------------------------------------------------
# Canonical quantities to verify consistency on
canonical_quantities = [
    # (pattern_for_value_with_label, human_label)
    (r'₹\s*7\.[0-9]\s*[Cc]rore', 'portfolio total'),
    (r'₹\s*7[,.]8[0-9]?\s*[Cc]rore', 'portfolio total'),
    (r'₹\s*3[,.]?86[,.]?648', 'LTCG carry-forward'),
    (r'₹\s*2[,.]?885\.4[0-9]', 'BSE avg cost'),
    (r'(\d+\.\d)x?\s+advantage', '3.8x advantage phrasing'),
    (r'1,?11[0-9]\s+orders', '1,113 orders friction'),
    (r'0\.0?33\s*%', '0.33% vs 0.033% friction rate'),
]

d11_hits = []
# Check BSE avg cost — should be ₹2,885.44 per project canonical
bse_variants = set()
for m in re.finditer(r'₹\s*2[,.]?88[0-9]\.[0-9]+', text):
    bse_variants.add(m.group(0))
if len(bse_variants) > 1:
    d11_hits.append(('BSE avg cost variants', list(bse_variants)))

# Check portfolio total — should consistently be ₹7.8 Crore per project canonical
# v3.4: Calculation-context guard. ₹ amounts appearing inside "SIP × years @ CAGR"
# arithmetic clauses are calculation outcomes, not portfolio-total claims.
# v3.3 false-positive: POST p1204 "₹12,000 SIP (15% of ₹82,000 EMI) compounding
# at 12% CAGR over 20 years grows to approximately ₹1.8 Crore" flagged ₹1.8 Crore
# as a portfolio-total variant. Not a drift — it's an arithmetic outcome.
CALC_CONTEXT_MARKERS = [
    r'CAGR',              # Compound annual growth rate calc
    r'compounding at',    # Explicit compounding calc
    r'grows to',          # Terminal-value calc
    r'accumulates',       # Step-up arithmetic
    r'produces approximately',  # Common phrasing for calc outcomes
    r'\d+\s*(?:years?|yrs?)',   # "over 20 years", "15 yr"
    r'\bSIP\b',           # SIP arithmetic
    r'step[- ]?up',       # Step-up progression
    r'in\s+\d+\s*years',  # "in 15 years"
    r'over\s+\d+\s*years', # "over 20 years"
]
CALC_CONTEXT_RE = re.compile('|'.join(f'({m})' for m in CALC_CONTEXT_MARKERS), re.IGNORECASE)

pf_variants = set()
for m in re.finditer(r'₹\s*(\d\.\d)\s*[Cc]rore', text):
    v = m.group(1)
    # Only track in portfolio context, not generic
    start = max(0, m.start() - 100)
    context = text[start:m.end()+50].lower()
    if 'portfolio' in context or 'family' in context or 'total' in context or 'author' in context:
        # v3.4: calculation-context guard — skip if context indicates arithmetic outcome
        calc_window = text[max(0, m.start() - 200):m.end() + 100]
        if CALC_CONTEXT_RE.search(calc_window):
            continue
        pf_variants.add(v)
if len(pf_variants) > 1:
    d11_hits.append(('portfolio total variants', sorted(pf_variants)))

# Check advantage multiplier — 3.8x should be canonical (per earlier anti-shortcuts discussion about 3.6x bug)
adv_variants = set()
for m in re.finditer(r'(\d\.\d)\s*x\s+advantage', text, re.I):
    adv_variants.add(m.group(1))
if len(adv_variants) > 1:
    d11_hits.append(('advantage multiplier variants', sorted(adv_variants)))

# Check friction rate — 0.33% canonical (NOT 0.033%)
if '0.033%' in text and '0.33%' in text:
    d11_hits.append(('friction rate', ['0.033%', '0.33%', 'both present — should be 0.33% only']))

# Check orders count — 1,113 canonical
order_variants = set()
for m in re.finditer(r'([12],\d{3})\s+orders', text):
    order_variants.add(m.group(1))
if len(order_variants) > 1:
    d11_hits.append(('orders count variants', sorted(order_variants)))

d11_status = 'PASS' if len(d11_hits) == 0 else 'FAIL'
if d11_hits:
    detector_fails.append(f"D11 numerical_inconsistency={len(d11_hits)}")
print(f"\nD11 Numerical consistency: {len(d11_hits)} inconsistencies {d11_status}")
for label, variants in d11_hits:
    print(f"    {label}: {variants}")

# ------------------------------------------------------------
# D12: Rupee/Dollar pairing (author's precisionist standard)
# CALIBRATED: Only flags canonical portfolio-specific amounts without $ equivalent.
# Skips: reader-example amounts, threshold phrases ("≥₹50 Lakh"), round numbers.
# Flags only when:
#   - ₹X is a specific decimal amount (e.g. ₹7,31,847) likely author-portfolio-specific
#   - OR ₹X is a specific Crore amount (≥₹1 Crore with decimal like ₹7.8 Crore)
#   - AND the same paragraph has no $ companion
#   - AND the paragraph is narrative prose (≥30 words)
# ------------------------------------------------------------
d12_hits = []
# Canonical author amount patterns (specific, non-generic)
canonical_rupee_patterns = [
    r'₹\s*\d+[,.]?\d+[,.]?\d{2,3}\b',  # specific amounts like ₹71,995 or ₹2,885.44 or ₹46,734
    r'₹\s*\d+\.\d\s*[Cc]rore\b',  # specific Crore amounts like ₹7.8 Crore (not "₹5 Crore" round)
]
for idx, p in enumerate(paras):
    if workbook_start is not None and idx >= workbook_start:
        continue
    if 'TYPESETTER' in p.text:
        continue
    t = p.text
    wc = len(t.split())
    if wc < 30:
        continue
    # Skip tables (paragraphs that are short label-value pairs)
    if t.count('|') >= 2:
        continue
    # Find specific ₹ amounts
    rupee_amounts = []
    for pat in canonical_rupee_patterns:
        for m in re.finditer(pat, t):
            rupee_amounts.append(m.group(0))
    if not rupee_amounts:
        continue
    # Check for $ companion in same paragraph
    dollar_amounts = re.findall(r'\$[\d,.]+\s*(?:million|USD|billion|M|B|K|k)?', t)
    if rupee_amounts and not dollar_amounts:
        d12_hits.append((idx, rupee_amounts[0], t[:80]))

d12_status = 'PASS' if len(d12_hits) == 0 else 'WARN' if len(d12_hits) < 30 else 'FAIL'
if len(d12_hits) >= 30:
    detector_fails.append(f"D12 missing_dollar_pair={len(d12_hits)}")
elif d12_hits:
    detector_warns.append(f"D12 missing_dollar_pair={len(d12_hits)}")
print(f"\nD12 Rupee/Dollar pairing (canonical amounts): {len(d12_hits)} {d12_status}")
for idx, amt, snippet in d12_hits[:10]:
    print(f"    [p{idx}] {amt}: {snippet}...")

# ------------------------------------------------------------
# D13: Date consistency — same event referenced with different date strings
# Focus: key events (March 23 relapse, FY26 numbers, etc.)
# ------------------------------------------------------------
d13_hits = []
# "March 23" should be consistent — some paragraphs have "March 23 2025", some "March 23 2026", some just "March 23"
march23_contexts = []
for idx, p in enumerate(paras):
    t = p.text
    if 'March 23' in t:
        # Extract the full date context
        for m in re.finditer(r'March\s+23(?:,?\s*(20\d{2}))?(?:rd)?', t):
            march23_contexts.append((idx, m.group(0), m.group(1) or 'year-unspecified'))

years_found = set(c[2] for c in march23_contexts)
if len(years_found) > 1:
    # If both unspecified and specified year exist, that's OK
    # But if both 2025 and 2026 show up, that's inconsistency
    specified_years = {y for y in years_found if y != 'year-unspecified'}
    if len(specified_years) > 1:
        d13_hits.append(('March 23 event year', sorted(specified_years)))

# FY26 should be consistent with March 2026 end
# Check for "FY25" vs "FY26" drift
fy_mentions = set(re.findall(r'\bFY(2[0-9])\b', text))

d13_status = 'PASS' if len(d13_hits) == 0 else 'WARN'
if d13_hits:
    detector_warns.append(f"D13 date_inconsistency={len(d13_hits)}")
print(f"\nD13 Date consistency: {len(d13_hits)} inconsistencies {d13_status}")
for label, variants in d13_hits:
    print(f"    {label}: {variants}")

# ------------------------------------------------------------
# D14: Heading numbering continuity
# Checks: Chapter N gaps (missing chapters), sub-chapter placement order
# ------------------------------------------------------------
# Collect all chapter numbers with heading style
chapter_nums_ordered = []  # list of (paragraph_idx, chapter_label)
for idx, p in enumerate(paras):
    if p.style and 'Heading' in p.style.name:
        m = re.match(r'^CHAPTER\s+(\d+[A-Z]?)', p.text.strip())
        if m:
            chapter_nums_ordered.append((idx, m.group(1)))

# Check for gaps in main numbering (exclude sub-chapters A/B)
main_chapter_nums = []
for idx, lbl in chapter_nums_ordered:
    m = re.match(r'^(\d+)([A-Z]?)$', lbl)
    if m and not m.group(2):  # main chapter only (no letter suffix)
        main_chapter_nums.append((idx, int(m.group(1)), lbl))

d14_hits = []
# Check gap sequence
if main_chapter_nums:
    expected_prev = None
    for idx, num, lbl in main_chapter_nums:
        if expected_prev is not None and num != expected_prev + 1 and num != expected_prev:
            # Gap detected (but allow duplicates like 0+0A)
            if num > expected_prev + 1:
                d14_hits.append(('gap', f'CHAPTER {expected_prev} → CHAPTER {num} (missing {list(range(expected_prev+1, num))})'))
        expected_prev = num

# Check sub-chapter placement (12A should come after 12, not after 13)
for i, (idx, lbl) in enumerate(chapter_nums_ordered[:-1]):
    m = re.match(r'^(\d+)([A-Z])$', lbl)
    if m:
        parent_num = int(m.group(1))
        # Check the preceding heading was the parent chapter
        if i > 0:
            prev_lbl = chapter_nums_ordered[i-1][1]
            prev_m = re.match(r'^(\d+)', prev_lbl)
            if prev_m and int(prev_m.group(1)) != parent_num:
                d14_hits.append(('misplaced_subchapter', f'{lbl} follows {prev_lbl} (expected after {parent_num})'))

d14_status = 'PASS' if len(d14_hits) == 0 else 'FAIL'
if d14_hits:
    detector_fails.append(f"D14 heading_continuity={len(d14_hits)}")
print(f"\nD14 Heading numbering continuity: {len(d14_hits)} issues {d14_status}")
for kind, detail in d14_hits[:10]:
    print(f"    {kind}: {detail}")

# ------------------------------------------------------------
# D15: Part endgame positioning
# Each "PART X" heading should be followed (within N paragraphs) by a Part intro,
# and a Part endgame/transition should appear at the END of the Part (before next PART)
# Flags: endgame at wrong location
# ------------------------------------------------------------
d15_hits = []
# Find all PART N headings
part_headings = []
for idx, p in enumerate(paras):
    if p.style and 'Heading' in p.style.name:
        m = re.match(r'^PART\s+([A-D])', p.text.strip())
        if m:
            part_headings.append((idx, m.group(1), p.text.strip()))

# For each Part, check if endgame appears in last 5% of the Part OR first 10% (misplaced)
for i, (idx, letter, heading) in enumerate(part_headings):
    next_part_idx = part_headings[i+1][0] if i+1 < len(part_headings) else len(paras)
    part_span = next_part_idx - idx
    if part_span < 20:
        continue
    # Search for endgame within the first 10% — that's misplaced
    first_10pct_end = idx + max(5, part_span // 10)
    last_10pct_start = next_part_idx - max(5, part_span // 10)
    for j in range(idx + 1, min(first_10pct_end, len(paras))):
        t = paras[j].text.lower()
        if 'endgame' in t or 'part closure' in t or ('closes' in t and 'part' in t):
            d15_hits.append((j, f'PART {letter} endgame at start (p{j}, {(j-idx)*100//part_span}% into Part)'))
            break

d15_status = 'PASS' if len(d15_hits) == 0 else 'WARN'
if d15_hits:
    detector_warns.append(f"D15 part_endgame={len(d15_hits)}")
print(f"\nD15 Part endgame positioning: {len(d15_hits)} issues {d15_status}")
for idx, detail in d15_hits[:10]:
    print(f"    {detail}")

# ------------------------------------------------------------
# D16: Sub-chapter placement (already partially covered by D14)
# Extends D14 with a more thorough gap analysis, including expected sub-chapters
# ------------------------------------------------------------
d16_hits = []
expected_sub_chapters = {'0A', '12A', '12B', '21A', '24A', '28A', '28B', '32A', '32B', '34A', '43A'}
found_sub = set()
for idx, lbl in chapter_nums_ordered:
    if re.match(r'^\d+[A-Z]$', lbl):
        found_sub.add(lbl)
missing_sub = expected_sub_chapters - found_sub
if missing_sub:
    # Only flag if mainstream chapters exist — i.e. REWRITE may not have all sub-chapters by design
    # But POST should have all of them
    # Report as WARN since it depends on edition
    d16_hits.append(('missing_subchapters', sorted(missing_sub)))

# Extra check: sub-chapter ordering (42A before 43A, not after)
sub_chapter_order = [(idx, lbl) for idx, lbl in chapter_nums_ordered if re.match(r'^\d+[A-Z]$', lbl)]
for i in range(len(sub_chapter_order) - 1):
    idx_a, lbl_a = sub_chapter_order[i]
    idx_b, lbl_b = sub_chapter_order[i+1]
    num_a = int(re.match(r'(\d+)', lbl_a).group())
    num_b = int(re.match(r'(\d+)', lbl_b).group())
    if idx_b > idx_a and num_b < num_a:
        d16_hits.append(('out_of_order', f'{lbl_a} (p{idx_a}) followed by {lbl_b} (p{idx_b}) — {lbl_b} should appear earlier'))

d16_status = 'PASS' if len(d16_hits) == 0 else 'WARN'
if d16_hits:
    detector_warns.append(f"D16 subchapter={len(d16_hits)}")
print(f"\nD16 Sub-chapter placement: {len(d16_hits)} issues {d16_status}")
for kind, detail in d16_hits[:10]:
    print(f"    {kind}: {detail}")

# ------------------------------------------------------------
# D17: TOC-vs-body consistency
# The Table of Contents should list the same chapters found in the body.
# Compare TOC entries to actual chapter headings.
# ------------------------------------------------------------
d17_hits = []
# Find TOC section
toc_start = None
for idx, p in enumerate(paras):
    if 'TABLE OF CONTENTS' in p.text.upper() or 'TABLE of CONTENTS' in p.text:
        toc_start = idx
        break

if toc_start is not None:
    # Read TOC entries (next ~40 paragraphs after the TOC heading)
    toc_chapters = set()
    for j in range(toc_start + 1, min(toc_start + 100, len(paras))):
        t = paras[j].text
        # TOC entries look like "CHAPTER N:" or "Chapter N —" or "(Chapters N-M)"
        for m in re.finditer(r'[Cc]hapter[s]?\s+(\d+[A-Z]?)', t):
            toc_chapters.add(m.group(1))
        # Also "(Chapters N-M)" ranges
        for m in re.finditer(r'[Cc]hapters\s+(\d+)[–-](\d+[A-Z]?)', t):
            start = int(m.group(1))
            end_str = m.group(2)
            end_m = re.match(r'(\d+)([A-Z]?)', end_str)
            if end_m:
                for n in range(start, int(end_m.group(1)) + 1):
                    toc_chapters.add(str(n))
        # Stop when we hit a Heading 1 (back to body)
        if paras[j].style and 'Heading 1' in paras[j].style.name and j > toc_start + 3:
            break
    
    # Body chapters
    body_chapters = set(lbl for _, lbl in chapter_nums_ordered)
    # Chapters in TOC but not in body
    in_toc_not_body = toc_chapters - body_chapters
    # Chapters in body but not TOC
    in_body_not_toc = body_chapters - toc_chapters
    
    if in_toc_not_body:
        d17_hits.append(('in_TOC_missing_body', sorted(in_toc_not_body, key=lambda x: int(re.match(r'\d+', x).group()) if re.match(r'\d+', x) else 999)))
    if in_body_not_toc:
        d17_hits.append(('in_body_missing_TOC', sorted(in_body_not_toc, key=lambda x: int(re.match(r'\d+', x).group()) if re.match(r'\d+', x) else 999)))
else:
    d17_hits.append(('no_TOC_found', 'Could not locate TABLE OF CONTENTS section'))

d17_status = 'PASS' if len(d17_hits) == 0 else 'WARN'
if d17_hits:
    detector_warns.append(f"D17 toc_body_drift={len(d17_hits)}")
print(f"\nD17 TOC-vs-body consistency: {len(d17_hits)} issues {d17_status}")
for kind, detail in d17_hits[:10]:
    detail_str = str(detail)[:150]
    print(f"    {kind}: {detail_str}")

# ------------------------------------------------------------
# D18: Rule/Appendix/Protocol cross-references (beyond just Chapter refs)
# Extends 3B (which only checks Chapter refs).
# References to: Rule RN, Appendix X, Protocol Y, Framework Z — verify each target exists.
# ------------------------------------------------------------
d18_hits = []
# Find all "Rule RN" references
rule_refs = set(re.findall(r'Rule\s+R(\d+[A-Z]?)', text))
# Reuse rule_defs from D10
rule_refs_missing = rule_refs - rule_defs
if rule_refs_missing:
    d18_hits.append(('missing_rule_definitions', sorted(rule_refs_missing, key=lambda x: (int(re.match(r'\d+', x).group()), x))))

# Find all "Appendix X" references
appendix_refs = set(re.findall(r'Appendix\s+([A-Z])', text))
# Find actual Appendix headings
appendix_headings = set()
for p in paras:
    t = p.text.strip()
    m = re.match(r'^APPENDIX\s+([A-Z])', t)
    if m and p.style and 'Heading' in p.style.name:
        appendix_headings.add(m.group(1))
appendix_missing = appendix_refs - appendix_headings
if appendix_missing:
    d18_hits.append(('missing_appendices', sorted(appendix_missing)))

# Find all "Protocol X" references (named protocols like "War Dip Protocol")
protocol_refs = set()
for m in re.finditer(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+Protocol\b', text):
    name = m.group(1)
    # Filter to those that look like formal protocols (1-3 words)
    if len(name.split()) <= 3:
        protocol_refs.add(name + ' Protocol')

d18_status = 'PASS' if len(d18_hits) == 0 else 'FAIL'
if d18_hits:
    detector_fails.append(f"D18 crossref={len(d18_hits)}")
print(f"\nD18 Rule/Appendix/Protocol cross-references: {len(d18_hits)} issues {d18_status}")
for kind, detail in d18_hits[:5]:
    detail_str = str(detail)[:150]
    print(f"    {kind}: {detail_str}")
if len(protocol_refs) > 0:
    print(f"    INFO: {len(protocol_refs)} named protocols referenced: {sorted(list(protocol_refs))[:5]}...")

# ------------------------------------------------------------
# Detectors result summary
# ------------------------------------------------------------
det_total_issues = (len(d1_hits) + len(d2_hits) + len(d3_hits) +
                    len(hard_walls) + len(d5_hits) + len(d6_hits) +
                    len(d7_hits) + len(d8_hits) + len(d9_hits) +
                    len(d10_hits) + len(d11_hits) + len(d12_hits) +
                    len(d13_hits) + len(d14_hits) + len(d15_hits) +
                    len(d16_hits) + len(d17_hits) + len(d18_hits))
det_any_fail = len(detector_fails) > 0
print(f"\nDETECTORS RESULT: {'PASS' if not det_any_fail else 'FAIL'}")
print(f"  STRUCTURAL:")
print(f"    D1 truncations: {len(d1_hits)}")
print(f"    D2 editorial notes: {len(d2_hits)}")
print(f"    D3 placeholders: {len(d3_hits)}")
print(f"    D4 walls (hard/soft): {len(hard_walls)}/{len(soft_walls)}")
print(f"    D5 orphan headings: {len(d5_hits)}")
print(f"    D6 duplicates: {len(d6_hits)}")
print(f"    D7 mid-list truncations: {len(d7_hits)}")
print(f"    D8 bracket placeholders: {len(d8_hits)}")
print(f"  CONSISTENCY:")
print(f"    D9 adjacent repetition: {len(d9_hits)}")
print(f"    D10 rule coverage gaps: {len(d10_hits)}")
print(f"    D11 numerical inconsistency: {len(d11_hits)}")
print(f"    D12 missing $ pairing: {len(d12_hits)}")
print(f"    D13 date inconsistency: {len(d13_hits)}")
print(f"    D14 heading continuity: {len(d14_hits)}")
print(f"    D15 part endgame: {len(d15_hits)}")
print(f"    D16 subchapter placement: {len(d16_hits)}")
print(f"    D17 TOC-vs-body drift: {len(d17_hits)}")
print(f"    D18 rule/appendix refs: {len(d18_hits)}")
if detector_fails:
    print(f"  FAILING: {', '.join(detector_fails)}")
if detector_warns:
    print(f"  WARNING: {', '.join(detector_warns)}")

# ============================================================
# SUMMARY
# ============================================================
print(f"\n{'='*60}")
print(f"SUMMARY")
print(f"{'='*60}")
print(f"QC 22-point: {qc_result}")
print(f"Ghost computational: {'PASS' if not fails else 'FAIL'}")
print(f"Jargon: {jargon_defined}/{jargon_total} ({len(jargon_undefined)} household terms undefined)")
print(f"Signpost: {wtm} WHY THIS MATTERS, {od} ONE DECISION")
print(f"  ONE DECISION coverage: {'ALL chapters' if not chapters_missing_od else 'MISSING: ' + ', '.join(chapters_missing_od)}")
print(f"  Part endgame closures: {'ALL' if not parts_missing_endgame else 'MISSING: ' + ', '.join(parts_missing_endgame)}")
print(f"  Quick Install achieve: {'PASS' if has_qi_achieve else 'MISSING'}")
print(f"Typesetter: {len(ts_blocks)} blocks")
print(f"Cross-refs broken: {len(broken)}")
print(f"Copyright: {cp_pass}/{len(cp_checks)}")
print(f"Hyphenation: {'PASS' if lt_real+st_real==0 else 'FIX'}")
print(f"Contra defense: {contra_result}")
print(f"Detectors:")
print(f"  Structural: D1={len(d1_hits)} D2={len(d2_hits)} D3={len(d3_hits)} D4h={len(hard_walls)} D5={len(d5_hits)} D6={len(d6_hits)} D7={len(d7_hits)} D8={len(d8_hits)}")
print(f"  Consistency: D9={len(d9_hits)} D10={len(d10_hits)} D11={len(d11_hits)} D12={len(d12_hits)} D13={len(d13_hits)} D14={len(d14_hits)} D15={len(d15_hits)} D16={len(d16_hits)} D17={len(d17_hits)} D18={len(d18_hits)}")
print(f"  Result: {'PASS' if not det_any_fail else 'FAIL'}")
print(f"SHA: {sha}")
