#!/usr/bin/env python3
"""WOS MANDATORY VERIFICATION GATE - Run before any 'production-ready' declaration.
Usage: python3 WOS_MVG.py <path_to_docx>"""
import sys, os, re, zipfile, hashlib
from docx import Document

# ========================================================
# LAYER 8: INTEGRITY CHECKS (S35 — declared clean without these)
# ========================================================
def run_integrity_layer(doc, full):
    failures = []
    print(f"\n--- LAYER 8: CONTENT INTEGRITY (from 682-test report) ---")
    
    import re
    # Check 1: Known word stutters only (not sentence boundaries)
    known_stutters = ["Bucket Bucket", "Relapse relapse", " the the ", " and and ", " or or ", " but but ", " for for "]
    for pattern in known_stutters:
        if pattern in full:
            for i, p in enumerate(doc.paragraphs):
                if pattern in p.text:
                    failures.append(f"STUTTER [{i}]: '{pattern}'")
    if not any("STUTTER" in f for f in failures):
        print(f"  ✅ No word stutters")
    
    # Check 2: Typesetter instructions in body
    typesetter = [i for i, p in enumerate(doc.paragraphs) if "[TYPESETTER:" in p.text]
    if typesetter:
        failures.append(f"TYPESETTER: {len(typesetter)} [TYPESETTER:] blocks at paras {typesetter} — will print as-is")
    else:
        print(f"  ✅ No typesetter instructions")
    
    # Check 3: Bandhan table specifically (known issue from integrity report)
    for i, t in enumerate(doc.tables):
        table_text = " ".join(cell.text for row in t.rows for cell in row.cells)
        if "Bandhan" in table_text:
            total = sum(1 for row in t.rows for cell in row.cells)
            empty = sum(1 for row in t.rows for cell in row.cells if not cell.text.strip())
            if empty / total > 0.8:
                failures.append(f"EMPTY TABLE [{i}]: Bandhan table {empty}/{total} cells empty")
    if not any("EMPTY TABLE" in f for f in failures):
        print(f"  ✅ No empty content tables")
    
    return failures

def run_gate(path):
    failures, warnings = [], []
    with open(path,'rb') as f: sha=hashlib.sha256(f.read()).hexdigest()[:16]
    print(f"{'='*60}\nWOS MVG | {os.path.basename(path)} | SHA:{sha}\n{'='*60}")
    
    # L1: XML
    with zipfile.ZipFile(path,'r') as z: xml=z.read('word/document.xml').decode('utf-8',errors='replace')
    dc=len(re.findall(r'<w:del ',xml)); ic=len(re.findall(r'<w:ins ',xml))
    if dc+ic>0: failures.append(f"TRACKED CHANGES: {dc} del, {ic} ins")
    else: print("  ✅ Tracked changes: 0")
    if re.findall(r'w:name="[^"]*0\.033[^"]*"',xml): failures.append("ORPHANED BOOKMARKS with 0.033")
    else: print("  ✅ Bookmarks clean")
    
    # L2: Paragraphs
    doc=Document(path); full="\n".join(p.text for p in doc.paragraphs)
    words=sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
    print(f"  Words:{words:,} Paras:{len(doc.paragraphs)} Tables:{len(doc.tables)}")
    for pat,desc in [("approximately .","blank"),("represents a .","blank"),("generates  in","dblspc")]:
        if pat in full: failures.append(f"BLANK: '{pat}'")
    for pat in ["Crore0.33%","turnover.0.33%","range.0.33%","years.0.33%","THE 3.8x ADVANTAGE","0.33%3.6x"]:
        if pat in full: failures.append(f"JUNK: '{pat}'")
    if "0.033%" in full: failures.append("OLD RATE: 0.033%")
    if sum(1 for p in doc.paragraphs if "**" in p.text)>0: failures.append("MARKDOWN **")
    
    # L3: Styles
    styles={}
    for p in doc.paragraphs: s=p.style.name if p.style else "None"; styles[s]=styles.get(s,0)+1
    if styles.get("Heading 1",0)!=26: failures.append(f"H1={styles.get('Heading 1',0)} (exp 26)")
    if styles.get("Normal",0)>2: warnings.append(f"Normal={styles.get('Normal',0)}")
    
    # L4: Content
    for label,check in [("0.33%≥17",full.count("0.33%")>=17),("XIRR def","Extended Internal Rate" in full),
        ("STCG def","Short-Term Capital Gains" in full),("BSE 2885","2,885.44" in full),
        ("IBKR 609","$609.33" in full),("3.6x","3.6x" in full),("8.5x","8.5x" in full),
        ("₹1.43Cr","₹1.43 Crore" in full),("₹2.87Cr","₹2.87 Crore" in full),
        ("RD Ch20","constraint is not capital" in full),("RD Ch30","Multi-Family Office" in full),
        ("RD Ch34A","who inherits the system" in full),("Privacy","appear by real name" in full)]:
        if not check: failures.append(f"MISSING: {label}")
    
    # L5: Hard rules
    for label,check in [("Finology 30","Finology 30" not in full),("Marcellus","Marcellus CCP" not in full),
        ("Wisesheets","wisesheets" not in full.lower()),("ArjunA","Arjun Acharya" not in full),
        ("AppxH","Appendix H" not in full)]:
        if not check: failures.append(f"HARD RULE: {label}")
    
    # L6: Counts
    for t,c,tgt,chk in [("relapse session",sum(1 for p in doc.paragraphs if "relapse session" in p.text.lower()),"≤10",lambda c:c<=10),
        ("March 23",full.lower().count("march 23"),"≤20",lambda c:c<=20)]:
        if not chk(c): warnings.append(f"BOUNDARY: {t}={c} ({tgt})")
    
    # L8: Integrity
    l8 = run_integrity_layer(doc, full)
    failures.extend(l8)

    # Verdict
    print(f"\n{'='*60}")
    if failures:
        print(f"❌ FAILED — {len(failures)}F {len(warnings)}W")
        for f in failures: print(f"  ❌ {f}")
        for w in warnings: print(f"  ⚠️ {w}")
    elif warnings:
        print(f"⚠️ PASS WITH {len(warnings)} WARNINGS")
        for w in warnings: print(f"  ⚠️ {w}")
    else: print("✅ ALL GATES PASS")
    print(f"SHA:{sha} | {words:,}w | {len(doc.paragraphs)}p | {len(doc.tables)}t")
    return len(failures)==0

if __name__=="__main__":
    if len(sys.argv)<2: print("Usage: python3 WOS_MVG.py <file.docx>"); sys.exit(1)
    sys.exit(0 if run_gate(sys.argv[1]) else 1)



# AC1 ENFORCEMENT: Cross-SOP consistency check
def cross_sop_verify(old_value, new_value, directory="."):
    """After changing old_value to new_value in any file,
    grep ALL files to ensure no stale references remain."""
    import glob
    stale = []
    for f in glob.glob(os.path.join(directory, "*.md")) + glob.glob(os.path.join(directory, "*.py")):
        with open(f, 'r', errors='replace') as fh:
            if old_value in fh.read():
                stale.append(os.path.basename(f))
    if stale:
        print(f"  AC1 VIOLATION: '{old_value}' still in: {stale}")
        return False
    print(f"  AC1 PASS: '{old_value}' propagated to '{new_value}' in all files")
    return True
